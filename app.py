import streamlit as st
from groq import Groq
from PIL import Image
import io
import base64
import json
import re
import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from datetime import datetime


def parse_recipe_json(response_text):
    text = response_text.strip()
    if text.startswith("```json"):
        text = text[7:].split("```")[0].strip()
    elif text.startswith("```"):
        text = text[3:].split("```")[0].strip()
    start = text.find('{')
    if start == -1:
        return None
    brace_count = 0
    end = start
    for i in range(start, len(text)):
        if text[i] == '{':
            brace_count += 1
        elif text[i] == '}':
            brace_count -= 1
            if brace_count == 0:
                end = i + 1
                break
    if brace_count != 0:
        end = text.rfind('}') + 1
    json_text = text[start:end]
    try:
        recipe = json.loads(json_text)
        if 'steps' in recipe:
            recipe['steps'] = [clean_step_numbering(s) for s in recipe['steps']]
        return recipe
    except json.JSONDecodeError:
        json_text = json_text.replace(',}', '}').replace(',]', ']')
        try:
            recipe = json.loads(json_text)
            if 'steps' in recipe:
                recipe['steps'] = [clean_step_numbering(s) for s in recipe['steps']]
            return recipe
        except:
            return None


def clean_step_numbering(step_text):
    if not step_text:
        return step_text
    return re.sub(r'^(\d+)\.\s+\1\.\s+', r'\1. ', step_text.strip())


def strip_step_numbering(step_text):
    if not step_text:
        return step_text
    return re.sub(r'^\d+\.\s*', '', step_text.strip())


def extract_schema_recipe(soup):
    """Extract recipe from schema.org JSON-LD if available."""
    scripts = soup.find_all('script', type='application/ld+json')
    for script in scripts:
        try:
            content = script.string or ''
            if not content.strip():
                continue
            data = json.loads(content)
            candidates = []
            if isinstance(data, dict):
                if data.get('@type') == 'Recipe':
                    candidates.append(data)
                if '@graph' in data:
                    candidates.extend([item for item in data['@graph'] if item.get('@type') == 'Recipe'])
            elif isinstance(data, list):
                for item in data:
                    if isinstance(item, dict):
                        if item.get('@type') == 'Recipe':
                            candidates.append(item)
                        if '@graph' in item:
                            candidates.extend([g for g in item['@graph'] if g.get('@type') == 'Recipe'])
            
            for rec in candidates:
                title = rec.get('name', 'Untitled Recipe')
                ingredients = rec.get('recipeIngredient', [])
                if not isinstance(ingredients, list):
                    ingredients = []
                ingredients = [str(ing).strip() for ing in ingredients if ing]
                
                steps = []
                instructions = rec.get('recipeInstructions', [])
                if isinstance(instructions, list):
                    for step in instructions:
                        if isinstance(step, str):
                            steps.append(step.strip())
                        elif isinstance(step, dict):
                            text = step.get('text') or step.get('name') or ''
                            steps.append(text.strip())
                elif isinstance(instructions, str):
                    steps.append(instructions.strip())
                
                if title != 'Untitled Recipe' and ingredients:
                    return {
                        "title": title,
                        "ingredients": ingredients,
                        "steps": [s for s in steps if s]
                    }
        except:
            continue
    return None


# Groq client
try:
    api_key = None
    try:
        api_key = st.secrets["GROQ_API_KEY"]
    except:
        pass  # No secrets.toml, continue to other sources

    if not api_key:
        api_key = os.environ.get("GROQ_API_KEY")

    if not api_key:
        try:
            import config
            api_key = config.GROQ_API_KEY
        except:
            pass

    if not api_key:
        st.error("GROQ_API_KEY not found. Please set it in:")
        st.error("- config.py file")
        st.error("- GROQ_API_KEY environment variable")
        st.error("- Streamlit secrets (.streamlit/secrets.toml)")
        st.stop()

    client = Groq(api_key=api_key)
except Exception as e:
    st.error(f"Groq init failed: {e}")
    st.stop()


def generate_html_cookbook(recipes, title, selected_style):
    # Common parts (header HTML is shared)
    header_html = f"""
    <div class="header"><h1>{title}</h1><p>Our Family Recipes • {datetime.now().strftime('%B %Y')}</p></div>
    """

    # Three full CSS variants
    if selected_style == "Trendy Simple":
        css = """
        body { font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; background: #ffffff; color: #2c3e50; line-height: 1.8; }
        .header { text-align: center; padding: 60px 20px; background: linear-gradient(135deg, #eceff1 0%, #cfd8dc 100%); border-radius: 20px; margin-bottom: 50px; color: #2c3e50; }
        .recipe { background: #ffffff; padding: 40px; margin: 40px 0; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.05); border: 1px solid #ecf0f1; }
        .recipe-title { font-size: 34px; color: #2c3e50; border-bottom: 4px solid #3498db; padding-bottom: 12px; margin-bottom: 30px; }
        .section-title { font-size: 24px; color: #2980b9; margin: 30px 0 15px; font-weight: 600; }
        .ingredients { padding: 25px; background: #f0f8ff; border-radius: 10px; border-left: 6px solid #3498db; }
        .ingredient { margin: 12px 0; padding-left: 5px; }
        .steps { padding-left: 10px; }
        .step { margin: 25px 0; font-size: 17px; }
        .step-number { font-weight: bold; color: #3498db; margin-right: 12px; font-size: 1.3em; }
        @media print { body { background: white; margin: 0; } .recipe { box-shadow: none; border: none; page-break-inside: avoid; } }
        """

    elif selected_style == "Old School Farmhouse":
        css = """
        body { font-family: Georgia, 'Times New Roman', serif; max-width: 900px; margin: 40px auto; padding: 20px; background: #fffef5; color: #4e342e; line-height: 1.7; }
        .header { text-align: center; padding: 60px 20px; background: linear-gradient(135deg, #ffebcd 0%, #f5deb3 100%); border-radius: 20px; margin-bottom: 50px; }
        .recipe { background: white; padding: 45px; margin: 40px 0; border-radius: 15px; box-shadow: 0 6px 20px rgba(0,0,0,0.1); border: 1px solid #deb887; }
        .recipe-title { font-size: 36px; color: #8b4513; border-bottom: 4px double #d2691e; padding-bottom: 15px; }
        .section-title { font-size: 26px; color: #a0522d; margin: 35px 0 15px; }
        .ingredients { padding: 25px; background: #fffacd; border-radius: 12px; border: 2px dashed #deb887; }
        .ingredient { margin: 12px 0; padding-left: 10px; }
        .steps { padding-left: 10px; }
        .step { margin: 25px 0; }
        .step-number { font-weight: bold; color: #cd853f; margin-right: 12px; font-size: 1.2em; }
        @media print { body { background: white; } .recipe { box-shadow: none; page-break-inside: avoid; } }
        """

    else:  # The Food Lab
        css = """
        body { font-family: Arial, Helvetica, sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; background: #ffffff; color: #000000; line-height: 1.7; }
        .header { text-align: center; padding: 60px 20px; background: #ff6600; color: white; border-radius: 0; margin-bottom: 60px; }
        .recipe { background: #ffffff; padding: 40px; margin: 40px 0; border: 2px solid #333; border-radius: 0; }
        .recipe-title { font-size: 38px; color: #ff6600; border-bottom: 5px solid #ff6600; padding-bottom: 10px; letter-spacing: 1px; font-weight: bold; }
        .section-title { font-size: 24px; color: #333333; margin: 40px 0 20px; text-transform: uppercase; letter-spacing: 2px; border-bottom: 1px solid #ccc; padding-bottom: 8px; }
        .ingredients { padding: 20px; background: #f9f9f9; border-left: 8px solid #ff6600; }
        .ingredient { margin: 12px 0; font-size: 16px; }
        .steps { padding-left: 0; counter-reset: step-counter; }
        .step { margin: 30px 0; font-size: 17px; }
        .step-number { font-weight: bold; color: #ff6600; margin-right: 15px; font-size: 1.5em; }
        @media print { body { background: white; margin: 0; } .recipe { border: 1px solid #000; page-break-inside: avoid; } }
        """

    html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>{title}</title>
<style>{css}</style></head><body>
{header_html}
"""
    for recipe in recipes:
        html += f'<div class="recipe"><div class="recipe-title">{recipe.get("title", "Untitled")}</div>'
        html += '<div class="section-title">Ingredients</div><div class="ingredients">'
        for ing in recipe.get("ingredients", []):
            html += f'<div class="ingredient">• {ing}</div>'
        html += '</div><div class="section-title">Instructions</div><div class="steps">'
        for i, step in enumerate(recipe.get("steps", []), 1):
            clean = strip_step_numbering(step)
            html += f'<div class="step"><span class="step-number">{i}.</span>{clean}</div>'
        html += '</div></div>'
    html += '</body></html>'
    return html


def generate_docx_cookbook(recipes, title, one_per_page, selected_style):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Style-specific settings
    if selected_style == "Trendy Simple":
        normal_font = 'Arial'
        heading_font = 'Arial'
        heading_color = RGBColor(44, 62, 80)      # dark blue-gray
        accent_color = RGBColor(52, 152, 219)     # blue
    elif selected_style == "Old School Farmhouse":
        normal_font = 'Times New Roman'
        heading_font = 'Georgia'
        heading_color = RGBColor(139, 69, 19)     # saddlebrown
        accent_color = RGBColor(205, 133, 63)     # peru
    else:  # The Food Lab
        normal_font = 'Calibri'
        heading_font = 'Arial Black'
        heading_color = RGBColor(255, 102, 0)     # bright orange
        accent_color = RGBColor(255, 102, 0)

    # Apply to Word styles
    styles = doc.styles
    styles['Normal'].font.name = normal_font
    styles['Normal'].font.size = Pt(11)
    styles['Heading 1'].font.name = heading_font
    styles['Heading 1'].font.size = Pt(28)
    styles['Heading 1'].font.color.rgb = heading_color
    styles['Heading 2'].font.name = heading_font
    styles['Heading 2'].font.color.rgb = heading_color

    # Cover
    doc.add_heading(title, 0).alignment = 1
    p = doc.add_paragraph('Our Family Recipes\n\nCollected with love')
    p.alignment = 1
    doc.add_paragraph(f"{datetime.now().strftime('%B %Y')}").alignment = 1
    doc.add_page_break()

    # Index
    doc.add_heading('Recipes', level=1)
    for recipe in recipes:
        doc.add_paragraph(recipe.get('title', 'Untitled'), style='List Number')
    doc.add_page_break()

    # Recipes
    for recipe in recipes:
        doc.add_heading(recipe.get('title', 'Untitled'), level=1)
        doc.add_heading('Ingredients', level=2)
        for ing in recipe.get('ingredients', []):
            doc.add_paragraph(ing, style='List Bullet')
        doc.add_heading('Instructions', level=2)
        for i, step in enumerate(recipe.get('steps', []), 1):
            clean = strip_step_numbering(step)
            p = doc.add_paragraph()
            num_run = p.add_run(f"{i}. ")
            num_run.bold = True
            num_run.font.color.rgb = accent_color
            num_run.font.size = Pt(14)
            p.add_run(clean)

        if one_per_page:
            doc.add_page_break()

    return doc


# App UI
st.set_page_config(page_title="Family Cookbook", layout="centered")
st.title("🍲 Our Family Cookbook")

st.markdown("""
**Create a beautiful printable cookbook from recipes your family sends you.**

Family can send:
- Recipe website links (most common – just paste them!)
- Photos of handwritten recipes (save to your device and upload)
- Copied recipe text

The app extracts everything perfectly.  
The final Word file is fully editable – add photos, rearrange, export PDF for printing/binding.
""")

if "recipes" not in st.session_state:
    st.session_state.recipes = []
if "cookbook_title" not in st.session_state:
    st.session_state.cookbook_title = "Our Family Cookbook"

cookbook_title = st.text_input("Cookbook title", value=st.session_state.cookbook_title)
st.session_state.cookbook_title = cookbook_title

st.markdown("### Add recipes")
col1, col2 = st.columns(2)
with col1:
    uploaded_files = st.file_uploader("Upload recipe photos", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
with col2:
    recipe_links = st.text_area("Paste recipe website links (one per line)")

text_input = st.text_area("Or paste recipe text")

if st.button("✨ Extract & Add Recipes", type="primary", use_container_width=True):
    if not (uploaded_files or recipe_links.strip() or text_input.strip()):
        st.warning("Please add something first.")
    else:
        with st.spinner("Extracting recipes..."):
            new_recipes = []

            # Process uploaded files (photos)
            for file in uploaded_files or []:
                try:
                    img = Image.open(file).convert('RGB')
                    buffered = io.BytesIO()
                    img.save(buffered, format="JPEG", quality=95)
                    base64_img = base64.b64encode(buffered.getvalue()).decode()
                    response = client.chat.completions.create(
                        model="meta-llama/llama-4-scout-17b-16e-instruct",
                        messages=[{"role": "user", "content": [
                            {"type": "text", "text": "Extract the recipe EXACTLY as JSON only. Preserve all quantities, units, and original text. Keep steps exactly as shown. Format: {\"title\": \"...\", \"ingredients\": [\"...\"], \"steps\": [\"...\"]}"},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_img}"}}
                        ]}],
                        max_tokens=1500
                    )
                    recipe = parse_recipe_json(response.choices[0].message.content)
                    if recipe and recipe.get('title') and recipe.get('ingredients'):
                        new_recipes.append(recipe)
                except:
                    st.error(f"Failed to process {file.name}")

            # Process website links ONLY
            for url in [u.strip() for u in recipe_links.splitlines() if u.strip()]:
                try:
                    headers = {"User-Agent": "Mozilla/5.0 (compatible; FamilyCookbook/1.0)"}
                    resp = requests.get(url, headers=headers, timeout=15)
                    resp.raise_for_status()
                    soup = BeautifulSoup(resp.text, 'html.parser')
                    
                    # First: structured schema
                    recipe = extract_schema_recipe(soup)
                    if recipe:
                        new_recipes.append(recipe)
                        continue
                    
                    # Fallback: clean text + LLM
                    for tag in soup(["script", "style", "nav", "header", "footer", "aside", "advert"]):
                        tag.decompose()
                    page_text = soup.get_text(separator='\n', strip=True)
                    if len(page_text) > 15000:
                        page_text = page_text[:15000] + "\n\n... (truncated)"
                    
                    response = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role": "user", "content": f"""Extract ONLY the main recipe from this webpage text as valid JSON. 
Preserve exact ingredient lines and step phrasing. 
Format: {{"title": "Recipe Name", "ingredients": ["full line 1", "full line 2"], "steps": ["step 1", "step 2"]}}

Text:
{page_text}"""}],
                        max_tokens=1500
                    )
                    recipe = parse_recipe_json(response.choices[0].message.content)
                    if recipe and recipe.get('title') and recipe.get('ingredients'):
                        new_recipes.append(recipe)
                except Exception as e:
                    st.warning(f"Could not extract recipe from: {url}")

            # Process text input
            if text_input.strip():
                try:
                    response = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role": "user", "content": f"Extract the recipe EXACTLY as JSON only. Preserve everything. Format: {{\"title\": \"...\", \"ingredients\": [...], \"steps\": [...]}}\n\nText:\n{text_input}"}],
                        max_tokens=1500
                    )
                    recipe = parse_recipe_json(response.choices[0].message.content)
                    if recipe and recipe.get('title') and recipe.get('ingredients'):
                        new_recipes.append(recipe)
                except:
                    st.error("Text extraction failed")

            # Dedupe and add
            titles = {r['title'].lower() for r in st.session_state.recipes}
            added = 0
            for r in new_recipes:
                if r['title'].lower() not in titles:
                    st.session_state.recipes.append(r)
                    titles.add(r['title'].lower())
                    added += 1

            if added:
                st.success(f"Added {added} new recipe(s)!")
                st.session_state.recipes.sort(key=lambda r: r.get('title', '').lower())
                st.rerun()
            else:
                st.info("No new recipes added (possible duplicates or extraction issues).")

# Current collection display, backup, generate (unchanged from previous version)

if st.session_state.recipes:
    st.markdown(f"### Your recipes ({len(st.session_state.recipes)})")
    for i in range(len(st.session_state.recipes)-1, -1, -1):
        r = st.session_state.recipes[i]
        col1, col2 = st.columns([6,1])
        with col1:
            st.markdown(f"**{r.get('title', 'Untitled')}**")
        with col2:
            if st.button("Remove", key=f"rem_{i}"):
                st.session_state.recipes.pop(i)
                st.rerun()

    with st.expander("💾 Save/Load progress (for working over multiple days)"):
        col1, col2 = st.columns(2)
        with col1:
            json_data = json.dumps(st.session_state.recipes, indent=4, ensure_ascii=False)
            st.download_button(
                "Download backup",
                json_data,
                "cookbook_backup.json",
                "application/json"
            )
        with col2:
            backup_file = st.file_uploader("Upload backup", type="json", key="backup")
            if backup_file:
                try:
                    imported = json.load(backup_file)
                    st.session_state.recipes = imported
                    st.success("Progress restored!")
                    st.rerun()
                except:
                    st.error("Invalid backup file")

    st.markdown("### Generate your cookbook")
    col_style1, col_style2 = st.columns([1, 3])
    with col_style1:
        one_per_page = st.checkbox("One recipe per page (recommended for printing)", value=True)
    with col_style2:
        style = st.selectbox(
            "Choose cookbook aesthetic",
            ["Trendy Simple", "Old School Farmhouse", "The Food Lab"]
        )

    if st.button("📖 Create Cookbook", type="primary", use_container_width=True):
        if not st.session_state.recipes:
            st.error("No recipes yet!")
        else:
            # DOCX
            doc = generate_docx_cookbook(st.session_state.recipes, cookbook_title, one_per_page, style)
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)

            st.download_button(
                "📄 Download Editable Word File (.docx)",
                docx_io.getvalue(),
                f"{cookbook_title.replace(' ', '_')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

            # HTML preview
            html = generate_html_cookbook(st.session_state.recipes, cookbook_title, style)
            st.markdown("#### 📱 Live Preview (great on phone too)")
            st.html(html)

            st.success("Cookbook ready! Open the .docx to customize and print ❤️")

st.caption("Made with love and AI • Simple, private, and free to use")
